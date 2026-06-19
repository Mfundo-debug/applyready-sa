import html
import io
import re
from datetime import date, timedelta
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

# =============================================================================
# ApplyReady SA - Application Pack Readiness Checker
# Version 2.5: rule-based, privacy-aware application fit engine and CV formaliser; no external AI, no document storage.
# =============================================================================

st.set_page_config(
    page_title="ApplyReady SA",
    page_icon="🟢",
    layout="wide",
    initial_sidebar_state="expanded",
)

PRIMARY = "#16A34A"
PRIMARY_DARK = "#08713A"
INK = "#0F172A"
MUTED = "#64748B"
BG = "#F6F8FB"
CARD = "#FFFFFF"
BORDER = "#E2E8F0"
SOFT_GREEN = "#ECFDF3"
AMBER = "#F59E0B"
RED = "#EF4444"
BLUE = "#2563EB"

st.markdown(
    f"""
    <style>
        :root {{
            --primary: {PRIMARY};
            --primary-dark: {PRIMARY_DARK};
            --ink: {INK};
            --muted: {MUTED};
            --bg: {BG};
            --card: {CARD};
            --border: {BORDER};
            --soft-green: {SOFT_GREEN};
        }}

        html, body, [class*="css"] {{
            font-family: "Inter", "Segoe UI", system-ui, -apple-system, BlinkMacSystemFont, sans-serif;
        }}

        .stApp {{
            background: radial-gradient(circle at top left, #EAFBF1 0, #F6F8FB 32%, #F6F8FB 100%);
            color: var(--ink);
        }}

        .block-container {{
            max-width: 1240px;
            padding-top: 1.25rem;
            padding-bottom: 3rem;
        }}

        section[data-testid="stSidebar"] {{
            background: linear-gradient(180deg, #0F172A 0%, #111827 100%);
            border-right: 1px solid rgba(255,255,255,0.08);
        }}

        section[data-testid="stSidebar"] * {{
            color: #E5E7EB !important;
        }}

        section[data-testid="stSidebar"] .stButton > button {{
            background: #22C55E !important;
            color: #052E16 !important;
            border: none !important;
            font-weight: 800 !important;
            border-radius: 14px !important;
        }}

        .hero {{
            position: relative;
            overflow: hidden;
            background: linear-gradient(135deg, #052E16 0%, #0B5F34 45%, #16A34A 100%);
            border-radius: 30px;
            padding: 2.25rem 2.25rem 2rem 2.25rem;
            color: white;
            box-shadow: 0 24px 70px rgba(15, 23, 42, 0.18);
            margin-bottom: 1.2rem;
        }}

        .hero:after {{
            content: "";
            position: absolute;
            width: 360px;
            height: 360px;
            right: -120px;
            top: -130px;
            background: rgba(255,255,255,0.16);
            border-radius: 999px;
        }}

        .hero-kicker {{
            display: inline-flex;
            align-items: center;
            gap: 0.5rem;
            background: rgba(255,255,255,0.14);
            border: 1px solid rgba(255,255,255,0.20);
            color: #DCFCE7;
            padding: 0.45rem 0.75rem;
            border-radius: 999px;
            font-size: 0.84rem;
            font-weight: 800;
            letter-spacing: 0.02em;
            text-transform: uppercase;
            margin-bottom: 1.1rem;
        }}

        .hero h1 {{
            margin: 0;
            font-size: clamp(2.35rem, 5vw, 4.7rem);
            line-height: 0.96;
            letter-spacing: -0.06em;
            font-weight: 950;
            color: white;
        }}

        .hero p {{
            max-width: 760px;
            margin: 1rem 0 0 0;
            color: #EAFBF1;
            font-size: 1.08rem;
            line-height: 1.7;
        }}

        .hero-grid {{
            display: grid;
            grid-template-columns: repeat(3, minmax(0, 1fr));
            gap: 0.9rem;
            margin-top: 1.35rem;
        }}

        .hero-stat {{
            background: rgba(255,255,255,0.12);
            border: 1px solid rgba(255,255,255,0.17);
            border-radius: 18px;
            padding: 0.95rem 1rem;
            backdrop-filter: blur(8px);
        }}

        .hero-stat strong {{
            display: block;
            font-size: 1.45rem;
            letter-spacing: -0.03em;
            color: white;
        }}

        .hero-stat span {{
            display: block;
            color: #BBF7D0;
            font-size: 0.86rem;
            margin-top: 0.15rem;
        }}

        .mini-card, .panel, .score-card, .fix-card {{
            background: rgba(255,255,255,0.92);
            border: 1px solid var(--border);
            border-radius: 22px;
            box-shadow: 0 12px 35px rgba(15,23,42,0.07);
        }}

        .mini-card {{
            padding: 1rem 1.05rem;
            min-height: 120px;
        }}

        .mini-card .icon {{
            width: 42px;
            height: 42px;
            border-radius: 14px;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            background: var(--soft-green);
            color: var(--primary-dark);
            font-weight: 900;
            margin-bottom: 0.65rem;
        }}

        .mini-card h3 {{
            margin: 0 0 0.3rem 0;
            color: var(--ink);
            font-size: 1.02rem;
        }}

        .mini-card p {{
            margin: 0;
            color: var(--muted);
            line-height: 1.55;
            font-size: 0.91rem;
        }}

        .panel {{
            padding: 1.15rem 1.2rem 1.25rem 1.2rem;
            margin-bottom: 1rem;
        }}

        .section-title {{
            display: flex;
            align-items: center;
            gap: 0.7rem;
            margin-bottom: 0.9rem;
        }}

        .section-title .num {{
            min-width: 34px;
            height: 34px;
            border-radius: 12px;
            background: #DCFCE7;
            color: #065F46;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            font-weight: 900;
        }}

        .section-title h2 {{
            margin: 0;
            color: var(--ink);
            font-size: 1.18rem;
            letter-spacing: -0.02em;
        }}

        .section-title p {{
            margin: 0.1rem 0 0 0;
            color: var(--muted);
            font-size: 0.9rem;
        }}

        .pill {{
            display: inline-flex;
            align-items: center;
            gap: 0.45rem;
            padding: 0.42rem 0.7rem;
            border-radius: 999px;
            font-size: 0.82rem;
            font-weight: 800;
            border: 1px solid var(--border);
            color: var(--ink);
            background: white;
            margin: 0.16rem 0.2rem 0.16rem 0;
        }}

        .pill.green {{ background: #DCFCE7; border-color: #BBF7D0; color: #065F46; }}
        .pill.amber {{ background: #FEF3C7; border-color: #FDE68A; color: #92400E; }}
        .pill.red {{ background: #FEE2E2; border-color: #FECACA; color: #991B1B; }}
        .pill.blue {{ background: #DBEAFE; border-color: #BFDBFE; color: #1E3A8A; }}

        .score-card {{
            padding: 1.1rem;
            min-height: 138px;
        }}

        .score-card .label {{
            font-size: 0.82rem;
            font-weight: 800;
            color: var(--muted);
            text-transform: uppercase;
            letter-spacing: 0.055em;
        }}

        .score-card .value {{
            font-size: 2.15rem;
            font-weight: 950;
            letter-spacing: -0.05em;
            color: var(--ink);
            margin-top: 0.15rem;
        }}

        .score-card .hint {{
            color: var(--muted);
            font-size: 0.88rem;
            margin-top: 0.3rem;
        }}

        .fix-card {{
            padding: 0.95rem 1rem;
            margin-bottom: 0.7rem;
            border-left: 5px solid #F59E0B;
            color: var(--ink);
        }}

        .fix-card strong {{ color: var(--ink); }}
        .fix-card span {{ color: var(--muted); }}

        .footer-note {{
            color: var(--muted);
            font-size: 0.88rem;
            padding: 1rem 0;
        }}

        /* Streamlit widget cleanup */
        label, .stTextInput label, .stTextArea label, .stSelectbox label, .stDateInput label, .stMultiSelect label {{
            color: #334155 !important;
            font-weight: 800 !important;
            font-size: 0.92rem !important;
        }}

        .stTextInput input, .stTextArea textarea, .stDateInput input {{
            background: #FFFFFF !important;
            color: #0F172A !important;
            border: 1px solid #CBD5E1 !important;
            border-radius: 14px !important;
        }}

        .stTextInput input:focus, .stTextArea textarea:focus {{
            border-color: #16A34A !important;
            box-shadow: 0 0 0 3px rgba(22,163,74,0.12) !important;
        }}

        div[data-baseweb="select"] > div {{
            background: #FFFFFF !important;
            color: #0F172A !important;
            border-color: #CBD5E1 !important;
            border-radius: 14px !important;
        }}

        .stTabs [data-baseweb="tab-list"] {{
            gap: 0.35rem;
            background: rgba(255,255,255,0.78);
            border: 1px solid var(--border);
            border-radius: 999px;
            padding: 0.35rem;
        }}

        .stTabs [data-baseweb="tab"] {{
            border-radius: 999px;
            padding: 0.55rem 1.0rem;
            font-weight: 850;
        }}

        .stTabs [aria-selected="true"] {{
            background: #DCFCE7 !important;
            color: #065F46 !important;
        }}

        div[data-testid="stDataFrame"] {{
            border-radius: 18px;
            overflow: hidden;
            border: 1px solid var(--border);
        }}

        div[data-testid="stProgress"] > div > div > div > div {{
            background: linear-gradient(90deg, #16A34A, #22C55E) !important;
        }}

        .stDownloadButton > button, .stButton > button {{
            border-radius: 14px !important;
            border: 1px solid #BBF7D0 !important;
            background: #16A34A !important;
            color: white !important;
            font-weight: 900 !important;
            box-shadow: 0 10px 20px rgba(22,163,74,0.18);
        }}

        .stDownloadButton > button:hover, .stButton > button:hover {{
            background: #08713A !important;
            border-color: #08713A !important;
        }}

        @media (max-width: 900px) {{
            .hero-grid {{ grid-template-columns: 1fr; }}
            .hero {{ padding: 1.5rem; border-radius: 22px; }}
        }}
    </style>
    """,
    unsafe_allow_html=True,
)


DOCUMENT_REQUIREMENTS: Dict[str, Dict[str, List[str]]] = {
    "Entry-level job": {
        "required": ["CV", "Certified ID copy", "Qualifications / certificates"],
        "recommended": ["Motivation letter", "Reference letter", "Academic record", "Proof of residence"],
    },
    "Internship": {
        "required": ["CV", "Certified ID copy", "Academic record", "Qualifications / certificates"],
        "recommended": ["Motivation letter", "Reference letter", "Proof of residence"],
    },
    "Learnership": {
        "required": ["CV", "Certified ID copy", "Matric certificate", "Proof of residence"],
        "recommended": ["Academic record", "Qualifications / certificates", "Motivation letter"],
    },
    "Bursary": {
        "required": [
            "Certified ID copy",
            "Academic record",
            "Proof of residence",
            "Proof of income / affidavit",
            "Proof of registration / acceptance letter",
        ],
        "recommended": ["Motivation letter", "CV", "Qualifications / certificates"],
    },
    "Graduate programme": {
        "required": ["CV", "Certified ID copy", "Academic record", "Qualifications / certificates"],
        "recommended": ["Motivation letter", "Reference letter", "Proof of residence"],
    },
}

DOCUMENT_DISPLAY_ORDER = [
    "Academic record",
    "CV",
    "Certified ID copy",
    "Matric certificate",
    "Motivation letter",
    "Proof of income / affidavit",
    "Proof of registration / acceptance letter",
    "Proof of residence",
    "Qualifications / certificates",
    "Reference letter",
]

# Keep the checklist clean and stable. This prevents duplicate-looking labels from
# appearing if a document name is repeated across opportunity types.
_available_docs = {
    doc.strip(): doc.strip()
    for req in DOCUMENT_REQUIREMENTS.values()
    for bucket in req.values()
    for doc in bucket
}
ALL_DOCUMENTS = [doc for doc in DOCUMENT_DISPLAY_ORDER if doc in _available_docs]

ACTION_WORDS = [
    "developed", "built", "created", "managed", "led", "analysed", "analyzed", "designed",
    "implemented", "supported", "coordinated", "improved", "assisted", "presented",
]

GENERIC_STOPWORDS = {
    "and", "the", "for", "with", "you", "your", "are", "our", "this", "that", "will", "must", "have",
    "has", "from", "into", "able", "about", "their", "they", "them", "can", "may", "any", "all",
    "job", "role", "candidate", "applicant", "application", "apply", "company", "programme", "program",
    "position", "opportunity", "requirements", "required", "preferred", "skills", "experience",
}

SKILL_KEYWORDS = [
    "python", "java", "sql", "excel", "power bi", "tableau", "data analysis", "machine learning",
    "communication", "teamwork", "problem solving", "leadership", "customer service", "admin",
    "administration", "sales", "marketing", "research", "report writing", "project management",
    "microsoft office", "word", "outlook", "presentation", "analytics", "accounting", "bookkeeping",
]

DEMO_VALUES = {
    "opportunity_type": "Internship",
    "company": "Ubuntu Analytics Lab",
    "role": "Junior Data Analyst Intern",
    "closing_date": date.today() + timedelta(days=14),
    "advert_text": """Ubuntu Analytics Lab is looking for a Junior Data Analyst Intern. Requirements include basic Python, Excel, SQL, communication skills, report writing, problem solving and teamwork. Applicants should submit a CV, certified ID copy, academic record, certificates and a short motivation letter.""",
    "advert_url": "",
    "full_name": "Taryn Michael",
    "email": "taryn@example.com",
    "qualification": "BSc Honours in Data Science",
    "field": "Data Science and Analytics",
    "location": "Kimberley, Northern Cape",
    "skills": "Python, Excel, SQL, data analysis, communication, teamwork, report writing",
    "experience": "Completed academic and personal projects involving exploratory data analysis, dashboards, predictive modelling and report writing.",
    "has_formal_experience": "No - I am a first-time applicant",
    "cv_text": """Taryn Michael\ntaryn@example.com | 071 234 5678 | Kimberley, Northern Cape\n\nEducation\nBSc Honours in Data Science, 2025\n\nSkills\nPython, Excel, SQL, data analysis, communication, teamwork, report writing\n\nProjects and Experience\nDeveloped data analysis notebooks using Python. Created dashboards and presented insights from student performance data. Assisted with academic projects involving machine learning and reporting.\n\nReferences available on request.""",
}


def init_state() -> None:
    defaults = {
        "opportunity_type": "Entry-level job",
        "company": "",
        "role": "",
        "closing_date": date.today() + timedelta(days=14),
        "advert_text": "",
        "advert_url": "",
        "full_name": "",
        "email": "",
        "qualification": "",
        "field": "",
        "location": "",
        "skills": "",
        "experience": "",
        "has_formal_experience": "No - I am a first-time applicant",
        "cv_text": "",
        "tracker": [],
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def load_demo() -> None:
    for key, value in DEMO_VALUES.items():
        st.session_state[key] = value
    for doc in ALL_DOCUMENTS:
        st.session_state[f"doc_{doc}"] = doc in [
            "CV", "Certified ID copy", "Academic record", "Qualifications / certificates", "Motivation letter"
        ]




def extract_text_from_uploaded_file(uploaded_file) -> Tuple[str, Optional[str]]:
    """Extract readable text from TXT, DOCX or PDF uploads without saving the file."""
    if uploaded_file is None:
        return "", "No file was uploaded."

    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    try:
        if file_name.endswith(".txt"):
            try:
                return file_bytes.decode("utf-8"), None
            except UnicodeDecodeError:
                return file_bytes.decode("latin-1", errors="ignore"), None

        if file_name.endswith(".docx"):
            document = Document(io.BytesIO(file_bytes))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            return "\n".join(paragraphs), None

        if file_name.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(file_bytes))
            page_text = []
            for page in reader.pages:
                extracted = page.extract_text() or ""
                if extracted.strip():
                    page_text.append(extracted.strip())
            text = "\n\n".join(page_text)
            if not text.strip():
                return "", "No readable text could be extracted. This may be a scanned PDF image."
            return text, None

        return "", "Unsupported file type. Please upload a PDF, DOCX or TXT file."
    except Exception as exc:
        return "", f"Could not extract text from {uploaded_file.name}: {exc}"


def fetch_job_advert_from_url(url: str) -> Tuple[str, Optional[str]]:
    """Fetch visible text from a job advert URL. Works best with public, static HTML pages."""
    url = (url or "").strip()
    if not url:
        return "", "Please paste a job advert link first."

    if not re.match(r"^https?://", url, flags=re.IGNORECASE):
        url = "https://" + url

    try:
        response = requests.get(
            url,
            timeout=12,
            headers={
                "User-Agent": "Mozilla/5.0 ApplyReadySA/1.0 (+https://streamlit.io)",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            },
        )
        response.raise_for_status()
    except Exception as exc:
        return "", f"Could not fetch the link. The site may block automated access or require login. Details: {exc}"

    content_type = response.headers.get("content-type", "").lower()
    if "pdf" in content_type:
        try:
            reader = PdfReader(io.BytesIO(response.content))
            text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
            return text.strip(), None if text.strip() else "The PDF link was fetched, but no readable text was found."
        except Exception as exc:
            return "", f"The link appears to be a PDF, but the app could not read it: {exc}"

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "nav", "footer", "header", "form"]):
        tag.decompose()

    pieces = []
    if soup.title and soup.title.string:
        pieces.append(soup.title.string.strip())

    for element in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        text = re.sub(r"\s+", " ", element.get_text(" ", strip=True)).strip()
        if len(text) >= 25:
            pieces.append(text)

    # De-duplicate while preserving order.
    seen = set()
    cleaned = []
    for piece in pieces:
        key = piece.lower()
        if key not in seen:
            seen.add(key)
            cleaned.append(piece)

    text = "\n".join(cleaned).strip()
    if not text:
        return "", "The page was fetched, but readable advert text was not found. Try uploading the advert PDF/DOCX or paste the text manually."

    return text[:15000], None


DOC_FILENAME_PATTERNS = {
    "CV": ["cv", "curriculum", "resume"],
    "Certified ID copy": ["id", "identity", "certified id"],
    "Academic record": ["academic", "record", "transcript", "results", "marks"],
    "Qualifications / certificates": ["certificate", "qualification", "degree", "diploma"],
    "Matric certificate": ["matric", "grade 12", "nsc"],
    "Proof of residence": ["residence", "address", "municipal", "utility"],
    "Proof of income / affidavit": ["income", "affidavit", "salary", "payslip"],
    "Proof of registration / acceptance letter": ["registration", "acceptance", "admission", "enrol", "enroll"],
    "Motivation letter": ["motivation", "cover letter", "cover_letter"],
    "Reference letter": ["reference", "recommendation", "referee"],
}


def infer_document_type_from_filename(file_name: str) -> Optional[str]:
    name = file_name.lower().replace("_", " ").replace("-", " ")
    for doc_type, patterns in DOC_FILENAME_PATTERNS.items():
        if any(pattern in name for pattern in patterns):
            return doc_type
    return None

def clean_name_for_file(name: str) -> str:
    name = name.strip() or "Applicant_Name"
    name = re.sub(r"[^A-Za-z0-9\s_-]", "", name)
    return re.sub(r"\s+", "_", name)


def extract_keywords(text: str, max_terms: int = 12) -> List[str]:
    text = text.lower()
    tokens = re.findall(r"[a-zA-Z][a-zA-Z+.#-]{2,}", text)
    counts: Dict[str, int] = {}
    for token in tokens:
        token = token.strip(".,;:()[]{}")
        if token in GENERIC_STOPWORDS or len(token) < 3:
            continue
        counts[token] = counts.get(token, 0) + 1

    for phrase in SKILL_KEYWORDS:
        if phrase in text:
            counts[phrase] = counts.get(phrase, 0) + 3

    ranked = sorted(counts.items(), key=lambda item: (-item[1], item[0]))
    return [word for word, _ in ranked[:max_terms]]


def unique_keep_order(values: List[str]) -> List[str]:
    seen = set()
    cleaned = []
    for value in values:
        item = " ".join(str(value).strip().split())
        if not item:
            continue
        key = item.lower()
        if key not in seen:
            seen.add(key)
            cleaned.append(item)
    return cleaned


def safe_html(value: Any) -> str:
    return html.escape(str(value), quote=True)


def extract_skills_from_text(text: str) -> List[str]:
    lower = (text or "").lower()
    found = []
    for skill in SKILL_KEYWORDS:
        pattern = r"\b" + re.escape(skill.lower()) + r"\b"
        if re.search(pattern, lower):
            found.append(skill.title() if skill.islower() else skill)

    # Capture common technology names that may appear with punctuation or uppercase.
    extra_patterns = {
        "C++": r"\bc\+\+\b",
        "C#": r"\bc#\b",
        "HTML": r"\bhtml\b",
        "CSS": r"\bcss\b",
        "JavaScript": r"\bjavascript\b",
        "R": r"\br programming\b|\br language\b",
        "SPSS": r"\bspss\b",
        "SAS": r"\bsas\b",
        "PowerPoint": r"\bpowerpoint\b",
    }
    for label, pattern in extra_patterns.items():
        if re.search(pattern, lower):
            found.append(label)
    return unique_keep_order(found)


def extract_required_documents_from_advert(text: str) -> List[str]:
    lower = (text or "").lower()
    found = []
    patterns = {
        "CV": ["cv", "curriculum vitae", "resume"],
        "Certified ID copy": ["certified id", "identity document", "id copy", "copy of id", "id document"],
        "Academic record": ["academic record", "transcript", "latest results", "statement of results"],
        "Qualifications / certificates": ["certificate", "qualification", "degree", "diploma"],
        "Matric certificate": ["matric certificate", "grade 12 certificate", "nsc certificate"],
        "Proof of residence": ["proof of residence", "proof of address", "municipal account"],
        "Proof of income / affidavit": ["proof of income", "affidavit", "payslip", "salary advice"],
        "Proof of registration / acceptance letter": ["proof of registration", "acceptance letter", "admission letter", "registration letter"],
        "Motivation letter": ["motivation letter", "cover letter", "letter of motivation"],
        "Reference letter": ["reference letter", "recommendation letter", "references"],
    }
    for doc, options in patterns.items():
        if any(option in lower for option in options):
            found.append(doc)
    return unique_keep_order(found)


def extract_qualification_requirements(text: str) -> List[str]:
    lower = (text or "").lower()
    requirements = []
    qualification_patterns = [
        ("Matric / Grade 12", r"matric|grade 12|nsc"),
        ("Certificate", r"certificate|higher certificate"),
        ("Diploma", r"diploma|national diploma"),
        ("Degree", r"degree|bachelor|bsc|ba|bcom|undergraduate"),
        ("Honours", r"honours|honors|postgraduate"),
        ("NQF level requirement", r"nqf\s*level\s*\d+|nqf\s*\d+"),
    ]
    for label, pattern in qualification_patterns:
        if re.search(pattern, lower):
            requirements.append(label)
    return unique_keep_order(requirements)


def extract_closing_date_mentions(text: str) -> List[str]:
    text = text or ""
    month_names = "january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec"
    patterns = [
        rf"\b\d{{1,2}}(?:st|nd|rd|th)?\s+(?:{month_names})\s+\d{{4}}\b",
        rf"\b(?:{month_names})\s+\d{{1,2}}(?:st|nd|rd|th)?[,]?\s+\d{{4}}\b",
        r"\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b",
        r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b",
    ]
    matches = []
    for pattern in patterns:
        matches.extend(re.findall(pattern, text, flags=re.IGNORECASE))
    return unique_keep_order([m if isinstance(m, str) else " ".join(m) for m in matches])[:5]


def infer_experience_level(text: str, opportunity_type: str) -> str:
    lower = (text or "").lower()
    if any(phrase in lower for phrase in ["no experience", "no prior experience", "training provided", "school leaver", "unemployed youth"]):
        return "First-time applicant friendly"
    if opportunity_type in ["Internship", "Learnership", "Bursary", "Graduate programme"]:
        return "Entry-level / early-career"
    years = re.findall(r"(\d+)\s*[-+]?\s*(?:year|years|yrs)", lower)
    if years:
        max_years = max(int(y) for y in years if y.isdigit())
        if max_years >= 3:
            return f"Experienced role ({max_years}+ years mentioned)"
        return f"Some experience preferred ({max_years}+ year(s) mentioned)"
    if any(word in lower for word in ["intern", "junior", "entry level", "entry-level", "graduate"]):
        return "Entry-level / early-career"
    return "Not clearly stated"


def smart_advert_analysis(advert_text: str, role: str, company: str, opportunity_type: str, closing_date: date) -> Tuple[Dict[str, Any], pd.DataFrame]:
    skills_found = extract_skills_from_text(advert_text)
    documents_found = extract_required_documents_from_advert(advert_text)
    qualifications_found = extract_qualification_requirements(advert_text)
    closing_mentions = extract_closing_date_mentions(advert_text)
    experience_level = infer_experience_level(advert_text, opportunity_type)
    keywords = extract_keywords(advert_text, max_terms=16)

    summary = {
        "Role/programme": role or "Use the title field or paste a clearer advert title",
        "Organisation": company or "Use the company/institution field",
        "Opportunity type": opportunity_type,
        "Experience level": experience_level,
        "Required qualification clues": ", ".join(qualifications_found) if qualifications_found else "Not clearly detected",
        "Skills detected": ", ".join(skills_found[:10]) if skills_found else "No common skills detected yet",
        "Documents detected in advert": ", ".join(documents_found) if documents_found else "No document list clearly detected",
        "Closing date field": closing_date.isoformat(),
        "Closing date text found": ", ".join(closing_mentions) if closing_mentions else "No date detected inside advert text",
        "Important keywords": ", ".join(keywords[:12]) if keywords else "Paste/upload/fetch an advert to detect keywords",
    }

    df = pd.DataFrame([{"Extracted item": key, "Result": value} for key, value in summary.items()])
    return summary, df


def application_fit_engine(cv_text: str, advert_text: str, qualification: str, skills: str, experience: str, selected_docs: List[str], opportunity_type: str, has_formal_experience: str) -> Dict[str, Any]:
    combined_cv = "\n".join([cv_text or "", qualification or "", skills or "", experience or ""]).lower()
    first_time = has_formal_experience.startswith("No")

    advert_skills = extract_skills_from_text(advert_text)
    if not advert_skills:
        advert_skills = extract_keywords(advert_text, max_terms=10)
    matched_skills = [skill for skill in advert_skills if skill.lower() in combined_cv]
    missing_skills = [skill for skill in advert_skills if skill.lower() not in combined_cv]
    skills_score = round((len(matched_skills) / len(advert_skills)) * 100) if advert_skills else 0

    advert_keywords = extract_keywords(advert_text, max_terms=16)
    matched_keywords = [kw for kw in advert_keywords if kw.lower() in combined_cv]
    missing_keywords = [kw for kw in advert_keywords if kw.lower() not in combined_cv]
    keyword_score = round((len(matched_keywords) / len(advert_keywords)) * 100) if advert_keywords else 0

    qualification_reqs = extract_qualification_requirements(advert_text)
    if not qualification_reqs:
        qualification_score = 70 if qualification.strip() else 0
    else:
        qualification_score = 100 if any(req.lower().split()[0] in combined_cv for req in qualification_reqs) else 45

    text_has_projects = any(term in combined_cv for term in ["project", "portfolio", "volunteer", "practical", "assignment", "dashboard", "model", "analysis", "leadership"])
    text_has_work = any(term in combined_cv for term in ["work experience", "employment", "worked", "intern", "assistant", "volunteer", "experience"])
    experience_level = infer_experience_level(advert_text, opportunity_type)
    if first_time and experience_level in ["First-time applicant friendly", "Entry-level / early-career"]:
        experience_score = 85 if text_has_projects else 65
    elif first_time:
        experience_score = 65 if text_has_projects else 45
    else:
        experience_score = 85 if text_has_work or text_has_projects else 45

    advert_docs = extract_required_documents_from_advert(advert_text)
    if advert_docs:
        docs_matched = [doc for doc in advert_docs if doc in selected_docs]
        document_fit_score = round((len(docs_matched) / len(advert_docs)) * 100)
    else:
        docs_matched = []
        document_fit_score = 70 if selected_docs else 0

    overall_fit = round(
        0.35 * skills_score
        + 0.20 * keyword_score
        + 0.15 * qualification_score
        + 0.15 * experience_score
        + 0.15 * document_fit_score
    )

    breakdown_df = pd.DataFrame([
        {"Fit area": "Skills match", "Score": skills_score, "Meaning": f"{len(matched_skills)} of {len(advert_skills)} advert skills reflected" if advert_skills else "No skills detected from advert"},
        {"Fit area": "Keyword alignment", "Score": keyword_score, "Meaning": f"{len(matched_keywords)} of {len(advert_keywords)} important advert keywords reflected" if advert_keywords else "No keywords detected from advert"},
        {"Fit area": "Qualification signal", "Score": qualification_score, "Meaning": ", ".join(qualification_reqs) if qualification_reqs else "No clear qualification requirement detected"},
        {"Fit area": "Experience/project evidence", "Score": experience_score, "Meaning": "First-time applicant mode active" if first_time else "Formal/project experience considered"},
        {"Fit area": "Advert document fit", "Score": document_fit_score, "Meaning": f"{len(docs_matched)} of {len(advert_docs)} advert documents ticked" if advert_docs else "No advert-specific document list detected"},
    ])

    recommendations = []
    if missing_skills:
        recommendations.append("Reflect relevant missing advert skills in the CV only if they are truthful: " + ", ".join(missing_skills[:8]) + ".")
    if missing_keywords:
        recommendations.append("Use the employer's language where truthful. Missing keywords include: " + ", ".join(missing_keywords[:8]) + ".")
    if first_time:
        recommendations.append("Use a Projects and Practical Experience section so the CV does not look empty even without formal work experience.")
    if advert_docs:
        missing_advert_docs = [doc for doc in advert_docs if doc not in selected_docs]
        if missing_advert_docs:
            recommendations.append("The advert appears to request these documents that are not ticked yet: " + ", ".join(missing_advert_docs) + ".")
    if overall_fit < 60:
        recommendations.append("Do not submit yet if you still have time. First improve the CV evidence and attach missing required documents.")

    return {
        "overall_fit": overall_fit,
        "breakdown_df": breakdown_df,
        "advert_skills": advert_skills,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "advert_keywords": advert_keywords,
        "matched_keywords": matched_keywords,
        "missing_keywords": missing_keywords,
        "qualification_reqs": qualification_reqs,
        "experience_level": experience_level,
        "first_time": first_time,
        "recommendations": recommendations,
    }


def fit_label(score: int) -> Tuple[str, str]:
    if score >= 80:
        return "Strong opportunity fit", "green"
    if score >= 65:
        return "Good fit with minor gaps", "green"
    if score >= 45:
        return "Possible fit, needs tailoring", "amber"
    return "Weak fit or incomplete evidence", "red"


def first_time_applicant_guidance(field: str, skills: str, experience: str, role: str) -> str:
    skill_items = [s.strip() for s in re.split(r",|\n", skills or "") if s.strip()]
    role_display = clean_sentence_fragment(role) or "the opportunity"
    practical_examples = clean_sentence_fragment(experience)
    if not practical_examples:
        practical_examples = "academic assignments, class projects, volunteering, short courses, leadership responsibilities or personal projects"

    return f"""Suggested CV section for a first-time applicant

Projects and Practical Experience
- Prepared for {role_display} by building practical exposure through {practical_examples}.
- Relevant skill areas: {', '.join(skill_items[:6]) if skill_items else 'communication, teamwork, problem solving and willingness to learn'}.
- I am still building formal workplace experience, but I can show reliability, learning ability and practical effort through projects, volunteering, coursework or community involvement.

Tip: Replace this with specific examples. For example, name one assignment, dashboard, report, volunteering task, leadership role or short course.
"""


def build_interview_pack(role: str, company: str, opportunity_type: str, matched_skills: List[str], missing_skills: List[str], first_time: bool) -> str:
    role_display = clean_sentence_fragment(role) or f"the {opportunity_type.lower()} opportunity"
    company_display = clean_sentence_fragment(company) or "the organisation"
    strengths = ", ".join(matched_skills[:4]) if matched_skills else "your strongest relevant skills"
    gaps = ", ".join(missing_skills[:3]) if missing_skills else "any areas you are still developing"

    first_time_note = "\nFirst-time applicant angle: be honest about limited formal experience, then redirect to projects, coursework, volunteering, reliability and willingness to learn.\n" if first_time else ""

    return f"""Interview Preparation Pack

Opportunity: {role_display} at {company_display}

Likely interview questions
1. Tell us about yourself and why you are interested in {role_display}.
   Suggested structure: qualification/background → relevant skills → why this opportunity interests you → what you can contribute.

2. Which skills make you suitable for this opportunity?
   Use examples linked to: {strengths}.

3. Tell us about a project, assignment, volunteering task or work experience you are proud of.
   Use STAR: Situation → Task → Action → Result.

4. What is one skill you are still improving?
   Mention the gap honestly, for example {gaps}, then explain what you are doing to improve.

5. Why should {company_display} consider your application?
   Focus on reliability, learning ability, communication, preparation and evidence from your CV.

Questions the applicant can ask
- What would success look like in the first three months?
- What training or support is available for new applicants?
- What qualities do successful candidates usually have in this role?

Documents to bring or keep ready
- CV
- Certified ID copy
- Academic record or certificates if requested
- Motivation letter or proof of application if requested
{first_time_note}
Confidence checklist
- I can explain my CV clearly.
- I can give at least one example of using my skills.
- I can explain why I want this opportunity.
- I know which documents were requested.
"""


def build_whatsapp_summary(role: str, company: str, closing_date: date, final_score: int, fit_score: int, missing_required: List[str], missing_skills: List[str], recommendations: List[str]) -> str:
    role_display = clean_sentence_fragment(role) or "Opportunity"
    company_display = clean_sentence_fragment(company) or "Organisation not provided"
    missing_doc_text = ", ".join(missing_required) if missing_required else "None detected"
    missing_skill_text = ", ".join(missing_skills[:5]) if missing_skills else "None detected"
    next_step = recommendations[0] if recommendations else "Review the official advert and submit before the closing date."
    return f"""ApplyReady SA Summary

Opportunity: {role_display}
Organisation: {company_display}
Closing date: {closing_date.isoformat()}
Readiness score: {final_score}/100
Opportunity fit score: {fit_score}/100
Missing required documents: {missing_doc_text}
Missing advert skills/keywords to check: {missing_skill_text}
Next step: {next_step}

Reminder: only add skills or experience that are truthful.
"""


def cv_checks(cv_text: str, advert_text: str) -> Tuple[pd.DataFrame, List[str], List[str], int, List[str]]:
    text = cv_text.strip()
    lower = text.lower()
    checks = []
    suggestions = []

    def add_check(area: str, passed: bool, note: str, weight: int):
        checks.append({"Area": area, "Status": "Ready" if passed else "Needs work", "Note": note, "Weight": weight})

    has_email = bool(re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text))
    has_phone = bool(re.search(r"(\+27|0)[0-9\s-]{8,}", text))
    contact_pass = has_email and has_phone
    add_check("Contact details", contact_pass, "Email and phone number found." if contact_pass else "Add a professional email address and reachable phone number.", 5)
    if not contact_pass:
        suggestions.append("Add clear contact details near the top of the CV: phone number, email address and location.")

    education_terms = ["education", "qualification", "degree", "diploma", "certificate", "matric", "university", "college", "school"]
    has_education = any(term in lower for term in education_terms)
    add_check("Education section", has_education, "Education/qualification information found." if has_education else "Add your qualification, institution and year completed or expected completion year.", 5)
    if not has_education:
        suggestions.append("Add an Education section with qualification, institution and dates.")

    has_skills = any(term in lower for term in ["skills", "technical skills", "computer skills", "competencies"])
    add_check("Skills section", has_skills, "Skills section found." if has_skills else "Add a Skills section with relevant technical and soft skills.", 5)
    if not has_skills:
        suggestions.append("Add a Skills section. Include only skills you can honestly defend in an interview.")

    experience_terms = ["experience", "employment", "work history", "volunteer", "project", "portfolio", "internship"]
    has_experience = any(term in lower for term in experience_terms)
    add_check("Experience or projects", has_experience, "Experience/project evidence found." if has_experience else "Add work experience, volunteering, academic projects or personal projects.", 5)
    if not has_experience:
        suggestions.append("Add an Experience or Projects section, even if the experience is academic, volunteer or personal project work.")

    has_dates = bool(re.search(r"\b(20\d{2}|19\d{2})\b", text))
    add_check("Dates", has_dates, "Dates found in the CV." if has_dates else "Add dates for education, work experience, projects and certificates.", 4)
    if not has_dates:
        suggestions.append("Add years or date ranges so the reader understands your timeline.")

    has_references = "reference" in lower or "references" in lower
    add_check("References", has_references, "Reference information found." if has_references else "Add references or write 'References available on request'.", 3)
    if not has_references:
        suggestions.append("Add a References section or state that references are available on request.")

    has_action_words = any(word in lower for word in ACTION_WORDS)
    add_check("Achievement language", has_action_words, "Action words found." if has_action_words else "Use action words such as developed, assisted, managed, created, analysed, supported or improved.", 4)
    if not has_action_words:
        suggestions.append("Rewrite bullet points with action words, for example: 'Developed', 'Assisted', 'Managed', 'Created', 'Analysed'.")

    word_count = len(re.findall(r"\w+", text))
    has_enough_detail = word_count >= 120
    add_check("CV detail", has_enough_detail, f"CV has about {word_count} words." if has_enough_detail else f"CV has about {word_count} words; it may be too short.", 4)
    if not has_enough_detail:
        suggestions.append("Add more detail. A very short CV usually does not show enough evidence of skills, education and experience.")

    advert_keywords = extract_keywords(advert_text)
    missing_keywords = [kw for kw in advert_keywords if kw.lower() not in lower]
    matched_keywords = [kw for kw in advert_keywords if kw.lower() in lower]
    alignment_pass = len(matched_keywords) >= max(2, min(5, len(advert_keywords) // 3)) if advert_keywords else False
    add_check("Advert keyword alignment", alignment_pass, f"Matched keywords: {', '.join(matched_keywords[:8]) or 'None found'}." if advert_keywords else "Paste the advert text to check keyword alignment.", 10)
    if missing_keywords:
        suggestions.append("Consider reflecting relevant advert keywords in your CV if they are truthful: " + ", ".join(missing_keywords[:8]) + ".")

    possible_points = sum(item["Weight"] for item in checks)
    earned_points = sum(item["Weight"] for item in checks if item["Status"] == "Ready")
    cv_score = round((earned_points / possible_points) * 30) if possible_points else 0
    return pd.DataFrame(checks), suggestions, advert_keywords, cv_score, matched_keywords


def score_documents(opportunity_type: str, selected_docs: List[str]) -> Tuple[pd.DataFrame, int, List[str]]:
    requirements = DOCUMENT_REQUIREMENTS[opportunity_type]
    required = requirements["required"]
    recommended = requirements["recommended"]
    rows = []
    missing_required = []

    for doc in required:
        has_doc = doc in selected_docs
        if not has_doc:
            missing_required.append(doc)
        rows.append({"Document": doc, "Importance": "Required", "Status": "Ready" if has_doc else "Missing"})

    for doc in recommended:
        has_doc = doc in selected_docs
        rows.append({"Document": doc, "Importance": "Recommended", "Status": "Ready" if has_doc else "Missing"})

    required_score = sum(1 for doc in required if doc in selected_docs) / len(required) if required else 0
    recommended_score = sum(1 for doc in recommended if doc in selected_docs) / len(recommended) if recommended else 0
    document_score = round((required_score * 30) + (recommended_score * 10))
    return pd.DataFrame(rows), document_score, missing_required


def score_profile(full_name: str, qualification: str, field: str, skills: str, experience: str, location: str) -> Tuple[int, List[str]]:
    checks = [
        (bool(full_name.strip()), "Add your full name."),
        (bool(qualification.strip()), "Add your highest qualification."),
        (bool(field.strip()), "Add your field of study or area of interest."),
        (len([s for s in re.split(r",|\n", skills) if s.strip()]) >= 3, "Add at least three relevant skills."),
        (len(experience.strip()) >= 40, "Add a short description of work, volunteer, academic or project experience."),
        (bool(location.strip()), "Add your town/city or preferred work location."),
    ]
    passed = sum(1 for ok, _ in checks if ok)
    score = round((passed / len(checks)) * 10)
    suggestions = [message for ok, message in checks if not ok]
    return score, suggestions


def score_communication(company: str, role: str, full_name: str, email: str) -> Tuple[int, List[str]]:
    checks = [
        (bool(company.strip()), "Add the company or institution name."),
        (bool(role.strip()), "Add the role, programme or bursary title."),
        (bool(full_name.strip()), "Add your full name for the email subject line."),
        (bool(re.search(r"[\w\.-]+@[\w\.-]+\.\w+", email)), "Add a valid email address."),
    ]
    passed = sum(1 for ok, _ in checks if ok)
    score = round((passed / len(checks)) * 10)
    suggestions = [message for ok, message in checks if not ok]
    return score, suggestions


def readiness_label(score: int) -> Tuple[str, str]:
    if score >= 80:
        return "Strong application pack", "green"
    if score >= 60:
        return "Almost ready", "amber"
    if score >= 40:
        return "Needs improvement", "amber"
    return "Not ready to submit", "red"


def clean_sentence_fragment(text_value: str) -> str:
    cleaned = " ".join((text_value or "").strip().split())
    return cleaned.rstrip(".")


def build_motivation_letter(full_name: str, company: str, role: str, opportunity_type: str, qualification: str, field: str, skills: str, experience: str, location: str, matched_skills: Optional[List[str]] = None, missing_skills: Optional[List[str]] = None, first_time_applicant: bool = False) -> str:
    skill_list = [s.strip() for s in re.split(r",|\n", skills) if s.strip()]
    matched_skills = matched_skills or []
    missing_skills = missing_skills or []
    strongest_skills = matched_skills[:5] if matched_skills else skill_list[:5]
    top_skills = ", ".join(strongest_skills) if strongest_skills else "relevant skills"

    company_display = clean_sentence_fragment(company) or "your organisation"
    role_display = clean_sentence_fragment(role) or f"the {opportunity_type.lower()} opportunity"
    name_display = clean_sentence_fragment(full_name) or "[Your Name]"
    qualification_display = clean_sentence_fragment(qualification) or "the qualification I have completed or am currently completing"
    field_display = clean_sentence_fragment(field) or "my field of study"
    location_display = clean_sentence_fragment(location)
    experience_display = clean_sentence_fragment(experience)

    location_sentence = f" I am based in {location_display}." if location_display else ""

    if experience_display and first_time_applicant:
        experience_sentence = f" Although I am still building formal work experience, I have developed practical exposure through {experience_display}."
    elif experience_display:
        experience_sentence = f" My experience includes {experience_display}."
    elif first_time_applicant:
        experience_sentence = " Although I am still building formal work experience, I am ready to demonstrate my reliability, willingness to learn and practical potential."
    else:
        experience_sentence = " I am eager to apply my knowledge, learn from experienced professionals and contribute positively."

    fit_sentence = ""
    if matched_skills:
        fit_sentence = f" The advert appears to value skills such as {', '.join(matched_skills[:4])}, which I have reflected in my application."
    elif missing_skills:
        fit_sentence = " I have reviewed the requirements carefully and I am willing to keep developing the skills needed for this opportunity."

    return f"""Dear Hiring Team,

I am writing to apply for {role_display} at {company_display}. I am interested in this opportunity because it aligns with my background in {field_display} and my goal of gaining practical experience through meaningful work.

My academic background includes {qualification_display}.{location_sentence} My relevant strengths include {top_skills}.{experience_sentence}{fit_sentence}

I believe I would be a suitable candidate because I am prepared, willing to learn, able to take responsibility and committed to contributing professionally. I would appreciate the opportunity to be considered for this {opportunity_type.lower()} and to demonstrate my potential.

Thank you for considering my application. I look forward to the possibility of hearing from you.

Kind regards,
{name_display}
"""


def build_email(full_name: str, company: str, role: str, opportunity_type: str, selected_docs: List[str]) -> Tuple[str, str]:
    name_display = full_name.strip() or "[Your Name]"
    role_display = role.strip() or f"{opportunity_type} opportunity"
    company_display = company.strip() or "your organisation"
    subject = f"Application for {role_display} – {name_display}"
    attachments = "\n".join([f"- {doc}" for doc in selected_docs]) if selected_docs else "- [Attach required documents]"
    body = f"""Dear Hiring Team,

I hope you are well.

Please find attached my application for {role_display} at {company_display}.

Attached documents:
{attachments}

Thank you for considering my application. I would appreciate the opportunity to be considered.

Kind regards,
{name_display}
"""
    return subject, body


def suggested_file_names(full_name: str, selected_docs: List[str]) -> pd.DataFrame:
    base = clean_name_for_file(full_name)
    mapping = {
        "CV": "CV.pdf",
        "Certified ID copy": "ID_Copy.pdf",
        "Academic record": "Academic_Record.pdf",
        "Qualifications / certificates": "Certificates.pdf",
        "Matric certificate": "Matric_Certificate.pdf",
        "Proof of residence": "Proof_of_Residence.pdf",
        "Proof of income / affidavit": "Proof_of_Income_or_Affidavit.pdf",
        "Proof of registration / acceptance letter": "Proof_of_Registration_or_Acceptance.pdf",
        "Motivation letter": "Motivation_Letter.pdf",
        "Reference letter": "Reference_Letter.pdf",
    }
    return pd.DataFrame(
        [{"Document": doc, "Suggested file name": f"{base}_{mapping.get(doc, doc.replace(' ', '_') + '.pdf')}"} for doc in selected_docs]
    )


def build_report(opportunity_type: str, company: str, role: str, closing_date: date, full_name: str, qualification: str, field: str, skills: str, experience: str, location: str, total_score: int, label: str, document_score: int, cv_score: int, profile_score: int, communication_score: int, doc_df: pd.DataFrame, cv_df: pd.DataFrame, suggestions: List[str], motivation_letter: str, email_subject: str, email_body: str, files_df: pd.DataFrame, fit_score: Optional[int] = None, fit_label_text: str = "", advert_summary_df: Optional[pd.DataFrame] = None, fit_breakdown_df: Optional[pd.DataFrame] = None, matched_skills: Optional[List[str]] = None, missing_skills: Optional[List[str]] = None, interview_pack: str = "", whatsapp_summary: str = "") -> str:
    buffer = io.StringIO()
    buffer.write("# ApplyReady SA - Application Readiness Report\n\n")
    buffer.write(f"Generated on: {date.today().isoformat()}\n\n")
    buffer.write("## Opportunity Details\n")
    buffer.write(f"- Opportunity type: {opportunity_type}\n")
    buffer.write(f"- Organisation: {company or '[Not provided]'}\n")
    buffer.write(f"- Role/programme: {role or '[Not provided]'}\n")
    buffer.write(f"- Closing date: {closing_date.isoformat()}\n\n")
    buffer.write("## Applicant Profile\n")
    buffer.write(f"- Full name: {full_name or '[Not provided]'}\n")
    buffer.write(f"- Highest qualification: {qualification or '[Not provided]'}\n")
    buffer.write(f"- Field of study: {field or '[Not provided]'}\n")
    buffer.write(f"- Skills: {skills or '[Not provided]'}\n")
    buffer.write(f"- Experience/projects: {experience or '[Not provided]'}\n")
    buffer.write(f"- Location: {location or '[Not provided]'}\n\n")
    buffer.write("## Readiness Score\n")
    buffer.write(f"- Total score: {total_score}/100\n")
    buffer.write(f"- Status: {label}\n")
    buffer.write(f"- Document readiness: {document_score}/40\n")
    buffer.write(f"- CV completeness and alignment: {cv_score}/30\n")
    buffer.write(f"- Applicant profile: {profile_score}/10\n")
    buffer.write(f"- Professional communication: {communication_score}/10\n\n")

    if fit_score is not None:
        buffer.write("## Version 2 Opportunity Fit Engine\n")
        buffer.write(f"- Opportunity fit score: {fit_score}/100\n")
        buffer.write(f"- Fit status: {fit_label_text}\n")
        if matched_skills:
            buffer.write(f"- Matched skills: {', '.join(matched_skills)}\n")
        if missing_skills:
            buffer.write(f"- Missing advert skills to review: {', '.join(missing_skills)}\n")
        buffer.write("\n")
        if advert_summary_df is not None and not advert_summary_df.empty:
            buffer.write("### Smart Advert Extraction\n")
            for _, row in advert_summary_df.iterrows():
                buffer.write(f"- {row['Extracted item']}: {row['Result']}\n")
            buffer.write("\n")
        if fit_breakdown_df is not None and not fit_breakdown_df.empty:
            buffer.write("### Fit Breakdown\n")
            for _, row in fit_breakdown_df.iterrows():
                buffer.write(f"- {row['Fit area']}: {row['Score']}/100 - {row['Meaning']}\n")
            buffer.write("\n")

    buffer.write("## Document Checklist\n")
    for _, row in doc_df.iterrows():
        buffer.write(f"- {row['Document']} ({row['Importance']}): {row['Status']}\n")
    buffer.write("\n## CV Checks\n")
    for _, row in cv_df.iterrows():
        buffer.write(f"- {row['Area']}: {row['Status']} - {row['Note']}\n")
    buffer.write("\n## Priority Fixes\n")
    if suggestions:
        for item in suggestions:
            buffer.write(f"- {item}\n")
    else:
        buffer.write("- No major issues detected. Review the application manually before submitting.\n")
    buffer.write("\n## Suggested File Names\n")
    if not files_df.empty:
        for _, row in files_df.iterrows():
            buffer.write(f"- {row['Document']}: {row['Suggested file name']}\n")
    else:
        buffer.write("- No documents selected.\n")
    buffer.write("\n## Motivation Letter Draft\n")
    buffer.write(motivation_letter.strip() + "\n\n")
    buffer.write("## Professional Email Draft\n")
    buffer.write(f"Subject: {email_subject}\n\n")
    buffer.write(email_body.strip() + "\n\n")
    if interview_pack:
        buffer.write("## Interview Preparation Pack\n")
        buffer.write(interview_pack.strip() + "\n\n")
    if whatsapp_summary:
        buffer.write("## WhatsApp-Friendly Summary\n")
        buffer.write(whatsapp_summary.strip() + "\n\n")
    buffer.write("## Important Note\n")
    buffer.write("This tool is a preparation assistant. Always verify the official advert requirements before submitting. Do not include false skills, qualifications or experience.\n")
    return buffer.getvalue()


def split_list_items(value: str) -> List[str]:
    """Split comma/newline separated text into clean, unique list items."""
    items = re.split(r",|\n|;", value or "")
    return unique_keep_order([item.strip() for item in items if item.strip()])


def extract_contact_hints(cv_text: str, email: str, location: str) -> Dict[str, str]:
    """Extract lightweight contact hints from the CV text without storing the file."""
    text = cv_text or ""
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    phone_match = re.search(r"(\+27|0)[0-9\s-]{8,}", text)
    return {
        "email": email.strip() or (email_match.group(0).strip() if email_match else "[Add professional email address]"),
        "phone": phone_match.group(0).strip() if phone_match else "[Add phone number]",
        "location": location.strip() or "[Add city/town]",
    }


def build_cv_polish_notes(cv_text: str, qualification: str, skills: str, experience: str, matched_skills: List[str], missing_skills: List[str], first_time: bool) -> List[str]:
    """Generate concise notes that help the applicant edit the formal CV draft."""
    text = cv_text or ""
    lower = text.lower()
    notes = []
    if not re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text):
        notes.append("Add a professional email address near the top of the CV.")
    if not re.search(r"(\+27|0)[0-9\s-]{8,}", text):
        notes.append("Add a reachable South African phone number.")
    if not qualification.strip() and not any(term in lower for term in ["matric", "diploma", "degree", "certificate", "education"]):
        notes.append("Add education details: qualification, institution and year completed or expected completion year.")
    if not split_list_items(skills) and not matched_skills:
        notes.append("Add a focused Skills section with technical and soft skills relevant to the opportunity.")
    if not experience.strip() and first_time:
        notes.append("Add at least one project, assignment, volunteering task, leadership role or short course under Practical Experience.")
    elif not experience.strip():
        notes.append("Add recent work experience, volunteering, projects or responsibilities with clear bullet points.")
    if not re.search(r"\b(20\d{2}|19\d{2})\b", text):
        notes.append("Add dates for studies, certificates, work experience, projects or volunteering.")
    if missing_skills:
        notes.append("Review missing advert skills and add only the ones the applicant can honestly explain: " + ", ".join(missing_skills[:6]) + ".")
    if not any(word in lower for word in ACTION_WORDS):
        notes.append("Use action words in bullet points, for example: developed, assisted, created, managed, supported or analysed.")
    if not notes:
        notes.append("The formal CV draft is ready for manual review. Check spelling, dates and truthfulness before submitting.")
    return unique_keep_order(notes)


def build_formal_cv_markdown(
    full_name: str,
    email: str,
    location: str,
    qualification: str,
    field: str,
    skills: str,
    experience: str,
    role: str,
    company: str,
    matched_skills: List[str],
    missing_skills: List[str],
    first_time: bool,
    cv_text: str,
) -> str:
    """Create a clean, formal, ATS-friendly CV draft using only supplied information."""
    name_display = clean_sentence_fragment(full_name) or "[Full Name]"
    contact = extract_contact_hints(cv_text, email, location)
    role_display = clean_sentence_fragment(role) or "the target opportunity"
    company_display = clean_sentence_fragment(company)
    field_display = clean_sentence_fragment(field) or "the relevant field"
    qualification_display = clean_sentence_fragment(qualification) or "[Add highest qualification]"
    experience_display = clean_sentence_fragment(experience)

    skill_items = split_list_items(skills)
    combined_skills = unique_keep_order(matched_skills + skill_items)
    if not combined_skills:
        combined_skills = ["Communication", "Teamwork", "Problem solving", "Willingness to learn"]

    skills_block = "\n".join([f"- {skill}" for skill in combined_skills[:12]])

    if first_time:
        profile = (
            f"Motivated {field_display} applicant with {qualification_display}. "
            f"Seeking an opportunity related to {role_display}"
            + (f" at {company_display}" if company_display else "")
            + ". Brings a willingness to learn, reliability and practical exposure through studies, projects, volunteering or personal development."
        )
        experience_heading = "Projects and Practical Experience"
        experience_intro = experience_display or "[Add academic assignments, personal projects, volunteering, leadership roles, short courses or community responsibilities.]"
        experience_bullets = [
            f"Built practical exposure through {experience_intro}",
            "Demonstrated learning ability, reliability and commitment through available academic or community-based experience.",
            f"Prepared to contribute to {role_display} by applying relevant skills and professional communication.",
        ]
    else:
        profile = (
            f"Professional applicant with a background in {field_display} and {qualification_display}. "
            f"Interested in {role_display}"
            + (f" at {company_display}" if company_display else "")
            + f". Offers strengths in {', '.join(combined_skills[:5])}."
        )
        experience_heading = "Work Experience and Projects"
        experience_intro = experience_display or "[Add recent work experience, internships, volunteering, projects or responsibilities.]"
        experience_bullets = [
            f"Experience includes {experience_intro}",
            "Used relevant skills to support tasks, solve problems and communicate results professionally.",
            "Able to work responsibly, learn quickly and contribute to team goals.",
        ]

    opportunity_alignment = []
    if matched_skills:
        opportunity_alignment.append("Skills already reflected for this opportunity: " + ", ".join(matched_skills[:8]) + ".")
    if missing_skills:
        opportunity_alignment.append("Skills to review and add only if truthful: " + ", ".join(missing_skills[:8]) + ".")
    if not opportunity_alignment:
        opportunity_alignment.append("Review the opportunity advert and tailor this CV to the most important requirements.")

    exp_block = "\n".join([f"- {bullet.rstrip('.')} .".replace(' .', '.') for bullet in experience_bullets])
    align_block = "\n".join([f"- {item}" for item in opportunity_alignment])

    return f"""# {name_display}

{contact['location']} | {contact['phone']} | {contact['email']}

## Professional Profile
{profile}

## Education
- {qualification_display}
- Field of study / interest: {field_display}
- Institution and year: [Add institution and year]

## Key Skills
{skills_block}

## {experience_heading}
{exp_block}

## Opportunity Alignment
{align_block}

## References
References available on request.

---
Editing reminder: replace bracketed placeholders, confirm all dates, and remove anything that is not truthful before submitting.
"""


def formal_cv_to_docx_bytes(cv_markdown: str, full_name: str) -> bytes:
    """Convert the formal CV markdown draft into a simple DOCX document."""
    document = Document()
    document.core_properties.title = "ApplyReady SA Formal CV Draft"
    document.core_properties.author = "ApplyReady SA"

    for raw_line in cv_markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("---"):
            continue
        else:
            document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def html_score_card(label: str, value: str, hint: str) -> str:
    return f"""
    <div class="score-card">
        <div class="label">{label}</div>
        <div class="value">{value}</div>
        <div class="hint">{hint}</div>
    </div>
    """


def section_title(num: str, title: str, subtitle: str) -> None:
    st.markdown(
        f"""
        <div class="section-title">
            <div class="num">{num}</div>
            <div>
                <h2>{title}</h2>
                <p>{subtitle}</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


init_state()

# Apply any widget state changes that were queued during the previous run.
# Streamlit does not allow changing a widget-backed session_state key
# after that widget has already been created in the same run.
if st.session_state.pop("mark_cv_doc_ready", False):
    st.session_state["doc_CV"] = True

# =============================================================================
# Sidebar
# =============================================================================

with st.sidebar:
    st.markdown("### 🟢 ApplyReady SA")
    st.caption("Application pack readiness checker")
    st.markdown("---")
    st.markdown("**Best use**")
    st.write("Paste/upload an advert, add applicant details, upload or tick available documents, then use the readiness report and CV draft before applying.")
    st.markdown("---")
    if st.button("Load polished demo data", use_container_width=True):
        load_demo()
        st.rerun()
    st.markdown("---")
    st.markdown("**Privacy-first beta**")
    st.write("No external AI API. Uploaded files are processed in memory only and are not stored by the app.")

# =============================================================================
# Header
# =============================================================================

st.markdown(
    """
    <div class="hero">
        <div class="hero-kicker">Version 2.5 beta • Built for job seekers, students and first-time applicants</div>
        <h1>ApplyReady SA</h1>
        <p>Prepare a complete, professional and opportunity-aligned application pack before you submit. Check missing documents, CV readiness, keyword alignment, email quality, file naming and create a formal CV draft in one place.</p>
        <div class="hero-grid">
            <div class="hero-stat"><strong>01</strong><span>Check application readiness</span></div>
            <div class="hero-stat"><strong>02</strong><span>Fix missing documents and weak CV signals</span></div>
            <div class="hero-stat"><strong>03</strong><span>Generate a report and formal CV draft</span></div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

c1, c2, c3 = st.columns(3)
with c1:
    st.markdown("""<div class="mini-card"><div class="icon">✓</div><h3>Readiness score</h3><p>Scores the application pack out of 100 using document, CV, profile and communication checks.</p></div>""", unsafe_allow_html=True)
with c2:
    st.markdown("""<div class="mini-card"><div class="icon">↗</div><h3>Advert alignment</h3><p>Accepts pasted adverts, uploaded advert files or public job-description links, then checks whether the CV reflects relevant skills honestly.</p></div>""", unsafe_allow_html=True)
with c3:
    st.markdown("""<div class="mini-card"><div class="icon">✦</div><h3>Professional pack</h3><p>Supports document upload, CV text extraction, formal CV drafting, motivation-letter drafts, email drafts, file-name suggestions and a downloadable report.</p></div>""", unsafe_allow_html=True)

st.write("")
setup_tab, docs_tab, fit_tab, cv_formaliser_tab, results_tab, drafts_tab, tracker_tab = st.tabs(
    ["1. Setup", "2. Documents & CV", "3. Fit Engine", "4. CV Formaliser", "5. Results", "6. Drafts & pack", "7. Tracker"]
)

# =============================================================================
# Inputs
# =============================================================================

with setup_tab:
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    section_title("1", "Opportunity details", "Capture the opportunity and the official requirements from the advert.")
    col1, col2, col3 = st.columns([1.05, 1.1, 0.85])
    with col1:
        st.selectbox("Opportunity type", list(DOCUMENT_REQUIREMENTS.keys()), key="opportunity_type")
        st.text_input("Company / institution name", placeholder="Example: ABC Company", key="company")
    with col2:
        st.text_input("Role / programme / bursary title", placeholder="Example: Data Analyst Internship", key="role")
        st.date_input("Closing date", key="closing_date")
    with col3:
        days_left = (st.session_state.closing_date - date.today()).days
        if days_left < 0:
            st.markdown(f'<span class="pill red">Closing date passed {-days_left} day(s) ago</span>', unsafe_allow_html=True)
        elif days_left <= 3:
            st.markdown(f'<span class="pill amber">Urgent: {days_left} day(s) left</span>', unsafe_allow_html=True)
        else:
            st.markdown(f'<span class="pill green">{days_left} day(s) left</span>', unsafe_allow_html=True)
        st.caption("The app applies a small urgency penalty when the deadline is too close and required documents are missing.")

    source_left, source_right = st.columns(2)
    with source_left:
        advert_file = st.file_uploader(
            "Upload advert document",
            type=["pdf", "docx", "txt"],
            key="advert_upload",
            help="Optional. Upload a PDF, DOCX or TXT advert. The file is read in memory and not stored.",
        )
        if st.button("Extract advert from document", use_container_width=True):
            text, error = extract_text_from_uploaded_file(advert_file)
            if error:
                st.error(error)
            else:
                st.session_state.advert_text = text
                st.success("Advert text extracted from document.")
                st.rerun()

    with source_right:
        st.text_input(
            "Fetch advert from link",
            placeholder="Example: https://company.com/jobs/data-analyst-intern",
            key="advert_url",
        )
        if st.button("Fetch advert from link", use_container_width=True):
            with st.spinner("Trying to fetch the advert text..."):
                text, error = fetch_job_advert_from_url(st.session_state.advert_url)
            if error:
                st.warning(error)
            else:
                st.session_state.advert_text = text
                st.success("Advert text fetched from the link.")
                st.rerun()

    st.caption("You can paste manually, upload an advert document, or fetch text from a public job-description link. Some sites block automated link reading, so upload/paste remains the fallback.")
    st.text_area(
        "Opportunity advert / requirements text",
        height=190,
        placeholder="Paste the advert text here. Include duties, requirements, skills and documents requested.",
        key="advert_text",
    )
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="panel">', unsafe_allow_html=True)
    section_title("2", "Applicant profile", "Basic information used to assess readiness and generate application drafts.")
    a, b = st.columns(2)
    with a:
        st.text_input("Full name", placeholder="Example: Taryn Michael", key="full_name")
        st.text_input("Email address", placeholder="Example: name@email.com", key="email")
        st.text_input("Highest qualification", placeholder="Example: Diploma in IT / BSc Data Science / Matric", key="qualification")
        st.text_input("Field of study / interest", placeholder="Example: Data Science, IT, Business Admin", key="field")
    with b:
        st.text_input("Location", placeholder="Example: Kimberley, Northern Cape", key="location")
        st.selectbox(
            "Do you have formal work experience?",
            ["No - I am a first-time applicant", "Yes - I have formal work experience"],
            key="has_formal_experience",
            help="Version 2 uses this to avoid unfairly penalising first-time applicants. Projects, volunteering and coursework can still count as evidence.",
        )
        st.text_area("Skills", height=95, placeholder="Example: Python, Excel, SQL, communication, teamwork", key="skills")
        st.text_area("Experience / projects / volunteering", height=95, placeholder="Briefly describe work experience, volunteering, academic projects or personal projects.", key="experience")
    st.markdown('</div>', unsafe_allow_html=True)

with docs_tab:
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    section_title("3", "Document checklist", "Tick only the documents the applicant already has ready.")
    current_type = st.session_state.opportunity_type
    required_docs = DOCUMENT_REQUIREMENTS[current_type]["required"]
    recommended_docs = DOCUMENT_REQUIREMENTS[current_type]["recommended"]
    st.markdown("<span class='pill green'>Required</span> " + " ".join([f"<span class='pill'>{doc}</span>" for doc in required_docs]), unsafe_allow_html=True)
    st.markdown("<span class='pill blue'>Recommended</span> " + " ".join([f"<span class='pill'>{doc}</span>" for doc in recommended_docs]), unsafe_allow_html=True)
    st.write("")

    uploaded_supporting_docs = st.file_uploader(
        "Upload application documents",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="supporting_docs_upload",
        help="Optional. Upload available documents so the app can detect document types from file names. Files are not stored.",
    )
    if uploaded_supporting_docs:
        detected_rows = []
        detected_docs = []
        for uploaded_doc in uploaded_supporting_docs:
            detected_type = infer_document_type_from_filename(uploaded_doc.name)
            detected_rows.append({
                "Uploaded file": uploaded_doc.name,
                "Detected document type": detected_type or "Not detected - tick manually",
            })
            if detected_type:
                detected_docs.append(detected_type)

        st.dataframe(pd.DataFrame(detected_rows), use_container_width=True, hide_index=True)
        if st.button("Use uploaded file names to tick checklist", use_container_width=True):
            for detected_doc in set(detected_docs):
                st.session_state[f"doc_{detected_doc}"] = True
            if detected_docs:
                st.success("Checklist updated from uploaded file names.")
                st.rerun()
            else:
                st.warning("No document types were detected from the file names. Please tick the checklist manually.")

    doc_cols = st.columns(3)
    for index, doc in enumerate(ALL_DOCUMENTS):
        with doc_cols[index % 3]:
            st.checkbox(doc, value=st.session_state.get(f"doc_{doc}", False), key=f"doc_{doc}")
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="panel">', unsafe_allow_html=True)
    section_title("4", "CV text", "Paste CV text or upload a CV document so the app can check completeness and advert alignment.")
    cv_upload_left, cv_upload_right = st.columns([1, 1])
    with cv_upload_left:
        cv_file = st.file_uploader(
            "Upload CV document",
            type=["pdf", "docx", "txt"],
            key="cv_upload",
            help="Optional. Upload a PDF, DOCX or TXT CV. The text is extracted in memory and the file is not stored.",
        )
    with cv_upload_right:
        st.write("")
        st.write("")
        if st.button("Extract CV text from document", use_container_width=True):
            text, error = extract_text_from_uploaded_file(cv_file)
            if error:
                st.error(error)
            else:
                st.session_state.cv_text = text
                st.session_state["mark_cv_doc_ready"] = True
                st.success("CV text extracted and CV document marked as ready.")
                st.rerun()

    st.text_area(
        "CV text",
        height=320,
        placeholder="Paste the CV text here or extract it from an uploaded CV. The app checks for contact details, education, skills, experience/projects, dates, references and advert alignment.",
        key="cv_text",
    )
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Scoring
# =============================================================================

selected_docs = [doc for doc in ALL_DOCUMENTS if st.session_state.get(f"doc_{doc}", False)]
opportunity_type = st.session_state.opportunity_type
company = st.session_state.company
role = st.session_state.role
closing_date = st.session_state.closing_date
advert_text = st.session_state.advert_text
full_name = st.session_state.full_name
email = st.session_state.email
qualification = st.session_state.qualification
field = st.session_state.field
location = st.session_state.location
skills = st.session_state.skills
experience = st.session_state.experience
has_formal_experience = st.session_state.has_formal_experience
cv_text = st.session_state.cv_text

doc_df, document_score, missing_required = score_documents(opportunity_type, selected_docs)
cv_df, cv_suggestions, advert_keywords, cv_score, matched_keywords = cv_checks(cv_text, advert_text)
profile_score, profile_suggestions = score_profile(full_name, qualification, field, skills, experience, location)
communication_score, communication_suggestions = score_communication(company, role, full_name, email)

advert_summary, advert_summary_df = smart_advert_analysis(advert_text, role, company, opportunity_type, closing_date)
fit_results = application_fit_engine(cv_text, advert_text, qualification, skills, experience, selected_docs, opportunity_type, has_formal_experience)
fit_score_value = fit_results["overall_fit"]
fit_label_text, fit_label_class = fit_label(fit_score_value)
fit_breakdown_df = fit_results["breakdown_df"]

urgency_penalty = 0
if closing_date < date.today():
    urgency_penalty = 10
elif (closing_date - date.today()).days <= 3 and missing_required:
    urgency_penalty = 5

raw_total = document_score + cv_score + profile_score + communication_score + 10
final_score = max(0, min(100, raw_total - urgency_penalty))
label, label_class = readiness_label(final_score)

all_suggestions = []
all_suggestions.extend([f"Missing required document: {doc}." for doc in missing_required])
all_suggestions.extend(cv_suggestions)
all_suggestions.extend(profile_suggestions)
all_suggestions.extend(communication_suggestions)
if urgency_penalty == 10:
    all_suggestions.append("The closing date appears to have passed. Verify whether late applications are accepted before preparing further.")
elif urgency_penalty == 5:
    all_suggestions.append("The closing date is very close. Prioritise missing required documents before editing optional items.")

all_suggestions.extend(fit_results["recommendations"])
# Remove repeated recommendations while preserving priority order.
all_suggestions = unique_keep_order(all_suggestions)

motivation_letter = build_motivation_letter(
    full_name,
    company,
    role,
    opportunity_type,
    qualification,
    field,
    skills,
    experience,
    location,
    matched_skills=fit_results["matched_skills"],
    missing_skills=fit_results["missing_skills"],
    first_time_applicant=fit_results["first_time"],
)
email_subject, email_body = build_email(full_name, company, role, opportunity_type, selected_docs)
files_df = suggested_file_names(full_name, selected_docs)
interview_pack = build_interview_pack(
    role,
    company,
    opportunity_type,
    fit_results["matched_skills"],
    fit_results["missing_skills"],
    fit_results["first_time"],
)
whatsapp_summary = build_whatsapp_summary(
    role,
    company,
    closing_date,
    final_score,
    fit_score_value,
    missing_required,
    fit_results["missing_skills"],
    fit_results["recommendations"],
)

formal_cv_markdown = build_formal_cv_markdown(
    full_name=full_name,
    email=email,
    location=location,
    qualification=qualification,
    field=field,
    skills=skills,
    experience=experience,
    role=role,
    company=company,
    matched_skills=fit_results["matched_skills"],
    missing_skills=fit_results["missing_skills"],
    first_time=fit_results["first_time"],
    cv_text=cv_text,
)
cv_polish_notes = build_cv_polish_notes(
    cv_text=cv_text,
    qualification=qualification,
    skills=skills,
    experience=experience,
    matched_skills=fit_results["matched_skills"],
    missing_skills=fit_results["missing_skills"],
    first_time=fit_results["first_time"],
)
formal_cv_docx = formal_cv_to_docx_bytes(formal_cv_markdown, full_name)

# =============================================================================
# Fit Engine
# =============================================================================

with fit_tab:
    section_title("5", "Application Fit Engine", "Version 2 compares the CV, profile and documents against the opportunity requirements.")
    f1, f2, f3, f4 = st.columns(4)
    with f1:
        st.markdown(html_score_card("Opportunity fit", f"{fit_score_value}/100", fit_label_text), unsafe_allow_html=True)
        st.progress(fit_score_value / 100)
    with f2:
        st.markdown(html_score_card("Skills matched", f"{len(fit_results['matched_skills'])}/{len(fit_results['advert_skills'])}", "Advert skills reflected"), unsafe_allow_html=True)
    with f3:
        st.markdown(html_score_card("Keywords matched", f"{len(fit_results['matched_keywords'])}/{len(fit_results['advert_keywords'])}", "Important advert wording"), unsafe_allow_html=True)
    with f4:
        mode_text = "First-time mode" if fit_results["first_time"] else "Experience mode"
        st.markdown(html_score_card("Applicant mode", mode_text, fit_results["experience_level"]), unsafe_allow_html=True)

    st.markdown(f'<span class="pill {fit_label_class}">{safe_html(fit_label_text)}</span>', unsafe_allow_html=True)
    st.write("")

    left, right = st.columns([1.05, 0.95])
    with left:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Smart advert extraction")
        st.caption("Rule-based extraction from the pasted, uploaded or linked advert. Review manually before applying.")
        st.dataframe(advert_summary_df, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Fit score breakdown")
        st.dataframe(fit_breakdown_df, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with right:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Skills found in advert")
        if fit_results["advert_skills"]:
            st.markdown("".join([f'<span class="pill blue">{safe_html(skill)}</span>' for skill in fit_results["advert_skills"]]), unsafe_allow_html=True)
        else:
            st.info("No common skills were detected from the advert yet. Paste more of the advert requirements.")
        st.write("")
        st.subheader("Matched in CV/profile")
        if fit_results["matched_skills"]:
            st.markdown("".join([f'<span class="pill green">{safe_html(skill)}</span>' for skill in fit_results["matched_skills"]]), unsafe_allow_html=True)
        else:
            st.warning("No advert skills were clearly matched in the CV/profile text yet.")
        st.write("")
        st.subheader("Missing from CV/profile")
        if fit_results["missing_skills"]:
            st.markdown("".join([f'<span class="pill amber">{safe_html(skill)}</span>' for skill in fit_results["missing_skills"][:10]]), unsafe_allow_html=True)
            st.caption("Only add missing skills if the applicant genuinely has them.")
        else:
            st.success("No major skill gaps detected from the extracted advert skills.")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Fit recommendations")
        if fit_results["recommendations"]:
            for item in fit_results["recommendations"][:6]:
                st.markdown(f'<div class="fix-card"><strong>Action:</strong> <span>{safe_html(item)}</span></div>', unsafe_allow_html=True)
        else:
            st.success("No major fit recommendations detected. Still review the advert manually.")
        st.markdown('</div>', unsafe_allow_html=True)

    if fit_results["first_time"]:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("First-time applicant guidance")
        st.caption("This helps applicants show evidence even when they do not have formal work experience yet.")
        st.text_area("Suggested CV section", value=first_time_applicant_guidance(field, skills, experience, role), height=245)
        st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# CV Formaliser
# =============================================================================

with cv_formaliser_tab:
    section_title("6", "CV Formaliser", "Create a cleaner, more formal CV draft from the uploaded/pasted CV and applicant profile.")

    top_left, top_right = st.columns([1.05, 0.95])
    with top_left:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Formal CV draft")
        st.caption("This is a structured draft, not a final guarantee. Edit it before submitting and remove anything that is not accurate.")
        st.text_area("ATS-friendly formal CV draft", value=formal_cv_markdown, height=620)
        col_md, col_docx = st.columns(2)
        with col_md:
            st.download_button(
                "Download CV draft as Markdown",
                data=formal_cv_markdown.encode("utf-8"),
                file_name=f"{clean_name_for_file(full_name)}_Formal_CV_Draft.md",
                mime="text/markdown",
                use_container_width=True,
            )
        with col_docx:
            st.download_button(
                "Download CV draft as Word document",
                data=formal_cv_docx,
                file_name=f"{clean_name_for_file(full_name)}_Formal_CV_Draft.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        st.markdown('</div>', unsafe_allow_html=True)

    with top_right:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Polish checklist")
        st.caption("Fix these before sending the CV to an employer, bursary provider or programme coordinator.")
        for note in cv_polish_notes:
            st.markdown(f'<div class="fix-card"><strong>Check:</strong> <span>{safe_html(note)}</span></div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("What this feature does")
        st.markdown(
            """
            - Reorganises the applicant's supplied information into a formal CV structure.  
            - Uses the opportunity fit results to highlight relevant skills.  
            - Supports first-time applicants by using projects, coursework or volunteering as evidence.  
            - Does **not** invent experience, qualifications or skills.
            """
        )
        st.warning("Before submitting, replace bracketed placeholders, add institution names and dates, and verify every statement.")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Suggested next edits")
        if fit_results["missing_skills"]:
            st.write("Review these advert skills. Add only the ones the applicant genuinely has:")
            st.markdown("".join([f'<span class="pill amber">{safe_html(skill)}</span>' for skill in fit_results["missing_skills"][:8]]), unsafe_allow_html=True)
        else:
            st.success("The current CV/profile already reflects the main detected advert skills.")
        st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Results
# =============================================================================

with results_tab:
    section_title("7", "Application readiness results", "Review the score, missing items and key fixes before submitting.")
    r1, r2, r3, r4 = st.columns(4)
    with r1:
        st.markdown(html_score_card("Total score", f"{final_score}/100", label), unsafe_allow_html=True)
        st.progress(final_score / 100)
    with r2:
        st.markdown(html_score_card("Documents", f"{document_score}/40", f"{len(selected_docs)} selected"), unsafe_allow_html=True)
    with r3:
        st.markdown(html_score_card("CV readiness", f"{cv_score}/30", f"{len(matched_keywords)} keyword matches"), unsafe_allow_html=True)
    with r4:
        st.markdown(html_score_card("Profile + email", f"{profile_score + communication_score}/20", "Applicant and communication checks"), unsafe_allow_html=True)

    st.markdown(f'<span class="pill {label_class}">{label}</span>', unsafe_allow_html=True)
    if urgency_penalty:
        st.markdown(f'<span class="pill amber">Urgency penalty: -{urgency_penalty}</span>', unsafe_allow_html=True)

    st.write("")
    left, right = st.columns([1.1, 0.9])
    with left:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Score breakdown")
        breakdown_df = pd.DataFrame(
            [
                {"Area": "Document readiness", "Score": document_score, "Maximum": 40},
                {"Area": "CV completeness and alignment", "Score": cv_score, "Maximum": 30},
                {"Area": "Applicant profile completeness", "Score": profile_score, "Maximum": 10},
                {"Area": "Professional communication readiness", "Score": communication_score, "Maximum": 10},
                {"Area": "Report/email generation baseline", "Score": 10, "Maximum": 10},
                {"Area": "Urgency penalty", "Score": -urgency_penalty, "Maximum": 0},
            ]
        )
        st.dataframe(breakdown_df, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Document checklist result")
        st.dataframe(doc_df, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with right:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Priority fixes")
        if all_suggestions:
            for item in all_suggestions[:10]:
                st.markdown(f'<div class="fix-card"><strong>Fix:</strong> <span>{safe_html(item)}</span></div>', unsafe_allow_html=True)
        else:
            st.success("No major issues detected. Still verify the official advert before submitting.")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.subheader("Advert keywords detected")
        if advert_keywords:
            chips = "".join([f'<span class="pill blue">{safe_html(kw)}</span>' for kw in advert_keywords])
            st.markdown(chips, unsafe_allow_html=True)
        else:
            st.info("Paste the advert text to detect opportunity keywords.")
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="panel">', unsafe_allow_html=True)
    st.subheader("CV readiness checks")
    st.dataframe(cv_df.drop(columns=["Weight"]), use_container_width=True, hide_index=True)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="panel">', unsafe_allow_html=True)
    st.subheader("Suggested file names")
    if files_df.empty:
        st.info("Tick documents in the checklist to generate file name suggestions.")
    else:
        st.dataframe(files_df, use_container_width=True, hide_index=True)
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Drafts and report
# =============================================================================

with drafts_tab:
    d1, d2 = st.columns(2)
    with d1:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        section_title("8", "Motivation letter draft", "Template-based draft. Edit before sending and keep it truthful.")
        st.text_area("Draft motivation letter", value=motivation_letter, height=420)
        st.markdown('</div>', unsafe_allow_html=True)
    with d2:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        section_title("9", "Professional email draft", "Subject line, body and attachment list for the application email.")
        st.text_input("Email subject", value=email_subject)
        st.text_area("Email body", value=email_body, height=337)
        st.markdown('</div>', unsafe_allow_html=True)

    p1, p2 = st.columns(2)
    with p1:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        section_title("10", "Interview preparation pack", "Likely questions, answer structure and practical preparation notes.")
        st.text_area("Interview pack", value=interview_pack, height=370)
        st.markdown('</div>', unsafe_allow_html=True)
    with p2:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        section_title("11", "WhatsApp-friendly summary", "Short summary the applicant can save, copy or send to someone helping them.")
        st.text_area("WhatsApp summary", value=whatsapp_summary, height=370)
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="panel">', unsafe_allow_html=True)
    section_title("12", "Downloadable readiness report", "Save the full application preparation report as a Markdown file.")
    report = build_report(
        opportunity_type=opportunity_type,
        company=company,
        role=role,
        closing_date=closing_date,
        full_name=full_name,
        qualification=qualification,
        field=field,
        skills=skills,
        experience=experience,
        location=location,
        total_score=final_score,
        label=label,
        document_score=document_score,
        cv_score=cv_score,
        profile_score=profile_score,
        communication_score=communication_score,
        doc_df=doc_df,
        cv_df=cv_df,
        suggestions=all_suggestions,
        motivation_letter=motivation_letter,
        email_subject=email_subject,
        email_body=email_body,
        files_df=files_df,
        fit_score=fit_score_value,
        fit_label_text=fit_label_text,
        advert_summary_df=advert_summary_df,
        fit_breakdown_df=fit_breakdown_df,
        matched_skills=fit_results["matched_skills"],
        missing_skills=fit_results["missing_skills"],
        interview_pack=interview_pack,
        whatsapp_summary=whatsapp_summary,
    )
    st.text_area("Report preview", value=report, height=360)
    st.download_button(
        "Download readiness report",
        data=report.encode("utf-8"),
        file_name=f"{clean_name_for_file(full_name)}_ApplyReady_Report.md",
        mime="text/markdown",
        use_container_width=True,
    )
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Tracker
# =============================================================================

with tracker_tab:
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    section_title("13", "Application tracker", "Track submitted opportunities during the current browser session.")
    t1, t2, t3 = st.columns([1, 1, 1])
    with t1:
        status = st.selectbox("Application status", ["Preparing", "Submitted", "Waiting for response", "Interview", "Rejected", "Successful"])
    with t2:
        follow_up_date = st.date_input("Follow-up date", value=closing_date + timedelta(days=7), key="follow_up_date")
    with t3:
        st.write("")
        st.write("")
        if st.button("Add current application", use_container_width=True):
            st.session_state.tracker.append(
                {
                    "Opportunity": role or "Untitled opportunity",
                    "Organisation": company or "Not provided",
                    "Type": opportunity_type,
                    "Closing date": closing_date.isoformat(),
                    "Status": status,
                    "Follow-up date": follow_up_date.isoformat(),
                    "Readiness score": final_score,
                    "Fit score": fit_score_value,
                }
            )
            st.success("Application added to tracker.")

    tracker_df = pd.DataFrame(st.session_state.tracker)
    if tracker_df.empty:
        st.info("No applications added yet.")
    else:
        st.dataframe(tracker_df, use_container_width=True, hide_index=True)
        csv = tracker_df.to_csv(index=False).encode("utf-8")
        st.download_button("Download tracker CSV", data=csv, file_name="applyready_application_tracker.csv", mime="text/csv", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="footer-note">ApplyReady SA is a preparation tool. Always follow the official advert instructions and never include false information in an application.</div>', unsafe_allow_html=True)
